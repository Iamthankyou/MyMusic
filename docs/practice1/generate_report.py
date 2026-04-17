"""
Practice 1: Multi-Agent Ideation Pipeline
Word report — Focused on: why it's better, how to use it, real examples
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG_DIR = r'd:\PROJECT\MyMusic\docs\practice1'
IMG = {k: os.path.join(IMG_DIR, f'{k}.png') for k in
       ['pipeline_v5_maxpower','comparison_v5_evolution','agents_v5_detail',
        'techniques_v5','evidence_v5','setup_v5','sources_v5','series_v5']}

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
sec.left_margin = Cm(2); sec.right_margin = Cm(2)
sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)

# ── Helpers ──────────────────────────────────────────────────
def shade(cell, hx):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hx)
    pr.append(s)

def run(p, text, bold=False, italic=False, sz=10, color=None, font='Calibri'):
    r = p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(sz); r.font.name=font
    if color: r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return r

def lcell(cell, text):
    shade(cell, '2E4057')
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.bold=True; r.font.size=Pt(10); r.font.name='Calibri'
    r.font.color.rgb = RGBColor(255,255,255)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def vcell(cell, items):
    first = True
    for it in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph(); first=False
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run(p, it.get('text',''), bold=it.get('bold',False), sz=it.get('size',10),
            color=it.get('color'))

def heading(text, level=1):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold=True; r.font.name='Calibri'
    if level==1: r.font.size=Pt(14); r.font.color.rgb=RGBColor(0x2E,0x40,0x57)
    elif level==2: r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x48,0x6F,0x8C)
    else: r.font.size=Pt(10.5); r.font.color.rgb=RGBColor(0x1A,0x5C,0x8B)

def body(text, bullet=False, bold=False, sz=10):
    p = doc.add_paragraph()
    if bullet: p.style = doc.styles['List Bullet']
    run(p, text, bold=bold, sz=sz); return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.name='Courier New'; r.font.size=Pt(8.5)
    return p

def img(key, caption=None, w=6.0):
    path = IMG.get(key)
    if path and os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(w))
        if caption:
            c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = c.add_run(caption); r.italic=True; r.font.size=Pt(8.5)
            r.font.color.rgb=RGBColor(0x66,0x66,0x66); r.font.name='Calibri'

def info_table(rows):
    t = doc.add_table(rows=len(rows), cols=2); t.style='Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(l,c) in enumerate(rows):
        row=t.rows[i]; row.cells[0].width=Cm(4.5); row.cells[1].width=Cm(13)
        lcell(row.cells[0],l); vcell(row.cells[1],c)
    doc.add_paragraph()

def table(headers, data):
    nc=len(headers); t=doc.add_table(rows=1+len(data), cols=nc)
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],'2E4057')
        r=t.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(9); r.font.name='Calibri'
        r.font.color.rgb=RGBColor(255,255,255)
    for ri,rd in enumerate(data):
        bg = 'FFFFFF' if ri%2==0 else 'EEF2F7'
        for ci,ct in enumerate(rd):
            shade(t.rows[ri+1].cells[ci], bg)
            r=t.rows[ri+1].cells[ci].paragraphs[0].add_run(ct)
            r.font.size=Pt(9); r.font.name='Calibri'
    doc.add_paragraph()

def callout(text, bg='FFF8E1'):
    """Yellow callout box"""
    t = doc.add_table(rows=1, cols=1); t.style='Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade(t.rows[0].cells[0], bg)
    p = t.rows[0].cells[0].paragraphs[0]
    run(p, text, bold=True, sz=10)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI Practice — Best Practice Guide'); r.bold=True; r.font.size=Pt(20)
r.font.name='Calibri'; r.font.color.rgb=RGBColor(0x2E,0x40,0x57)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Practice 1 of 3'); r.font.size=Pt(14); r.font.name='Calibri'
r.font.color.rgb=RGBColor(0x48,0x6F,0x8C); r.bold=True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Multi-Agent Ideation Pipeline\nTurn 1 Sentence Into a Validated Product Concept\nUsing 8 AI Agents That Research the Internet For You')
r.font.size=Pt(12); r.font.name='Calibri'; r.font.color.rgb=RGBColor(0x48,0x6F,0x8C)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Submitted: 2026-04-16  |  Video Editor  |  Duy LQ')
r.font.size=Pt(10); r.font.name='Calibri'; r.font.color.rgb=RGBColor(0x99,0x99,0x99)

doc.add_paragraph()
img('series_v5', 'This is Practice 1 of a 3-part series: Ideation → Planning → Code', 5.5)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  1. THE PROBLEM — Why This Matters
# ══════════════════════════════════════════════════════════════
heading('1. The Problem: Why Most AI Brainstorming Is Useless', 1)

body('Most people brainstorm with AI like this:')
callout('💬 "Hey ChatGPT, I want to make a music app for indie artists. What do you think?"')

body('And they get back a wall of generic advice that SOUNDS good but:')
body('❌ Has no real data — the AI makes up market sizes and competitor names', bullet=True)
body('❌ Has no sources — you can\'t verify anything it says', bullet=True)
body('❌ Never challenges you — it agrees with everything you say', bullet=True)
body('❌ Misses competitors — it doesn\'t actually search the internet', bullet=True)
body('❌ Gives you confidence without evidence — the most dangerous outcome', bullet=True)

body('The result? You feel productive but you\'ve learned nothing real. '
     'You start building a product based on hallucinated data.', bold=True)

heading('What Does "Good" Brainstorming Look Like?', 2)
body('A proper product ideation should give you:')
body('✅ Real competitor data with URLs you can click and verify', bullet=True)
body('✅ Actual market size numbers from real research reports', bullet=True)
body('✅ A devil\'s advocate that challenges your assumptions', bullet=True)
body('✅ Alternative approaches you hadn\'t considered', bullet=True)
body('✅ An honest "should you build this?" verdict based on evidence', bullet=True)

body('This method gives you ALL of those — automatically.', bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  2. BEFORE vs AFTER — The Real Comparison
# ══════════════════════════════════════════════════════════════
heading('2. Traditional AI Brainstorming vs This Method', 1)

body('Let\'s compare with a real example. Imagine you have this idea:')
callout('💡 "I want to make a music streaming app focused on independent artists, like Spotify but for indie music."')

heading('What Happens with the Traditional Approach', 2)
body('You ask ChatGPT / Cursor / Claude in 1 chat session:', bold=True)

table(
    ['You Ask', 'AI Responds', 'The Problem'],
    [
        ['"What competitors exist?"',
         '"Some competitors include\nSpotify, SoundCloud, and\nBandcamp..."',
         'No URLs. No user counts.\nNo pricing data.\nMight not even be accurate.'],
        ['"What\'s the market size?"',
         '"The music streaming market\nis estimated at $30 billion..."',
         'Where did $30B come from?\nNo source. Could be made up.\nNo way to verify.'],
        ['"Is my idea good?"',
         '"Yes, this is a promising idea\nwith great potential..."',
         'Of course it says yes.\nIt\'s trained to be helpful.\nNo critical analysis.'],
        ['"What tech should I use?"',
         '"You could use React Native\nwith Firebase and Spotify API"',
         'Did it check Spotify API pricing?\nDid it check rate limits?\nNo.'],
    ]
)
body('Total time: ~15 minutes. Total verifiable facts: probably 0.', bold=True)

heading('What Happens with This Method', 2)
body('You type 1 sentence. Then 8 specialized AI agents work for you:', bold=True)

table(
    ['Agent Does', 'What You Get', 'Why It\'s Different'],
    [
        ['Agent 1 asks you 10-15\nstructured questions',
         'A proper Requirements Brief\nwith validated assumptions',
         'Forces you to think through WHO,\nWHAT, WHY, HOW, CONSTRAINTS.\nThen searches web to verify.'],
        ['Agent 2 searches the internet\nfor market data',
         'TAM: $28.6B (Grand View Research)\nSAM: $4.2B indie music\nwith source URLs',
         'Every number has a clickable\nsource URL. You can verify\neverything yourself.'],
        ['Agent 3 finds 5+ real\ncompetitors',
         'Bandcamp: 8M users, $201M rev\nDistroKid: $35.99/yr\n... with App Store links',
         'Real products with real URLs,\nreal pricing, real user quotes\nfrom Reddit/App Store reviews.'],
        ['Agent 4 searches for APIs\nand checks pricing',
         'Spotify API: free tier 100 req/min\nYouTube API: 10K units/day free\nFirebase: free up to 50K MAU',
         'Actual pricing pages checked.\nRate limits verified.\nNo hallucinated APIs.'],
        ['Agent 5 scores your idea\non 8 dimensions',
         'Feasibility: 4/5 — evidence: ...\nMarket demand: 3/5 — evidence: ...\nVerdict: CONDITIONAL GO',
         'Every score has 1 sentence\nof evidence. No free passes.\nDevil\'s advocate built in.'],
        ['Agent 6 designs the UX\nwith research',
         'Time-to-aha: 2 min (benchmark)\nOnboarding: 3 steps max\nHook model mapped',
         'Based on NNGroup/Baymard\nresearch, not guessing.\nCompetitor weakness analysis.'],
        ['Agent 7 identifies risks\nwith legal research',
         'Music licensing: CRITICAL risk\nGDPR compliance needed\nAPI dependency: HIGH risk',
         'Actually searches for\nregulations. Finds real legal\nrequirements with URLs.'],
        ['Agent 8 compiles everything\ninto 1 document',
         '13-Section Idea Specification\nFull evidence trail\nGO / PIVOT / KILL verdict',
         'Nothing made up. Every claim\ntraced to a source.\nConfidence score for spec.'],
    ]
)
body('Total time: ~20 minutes. Total verifiable facts: 30+.', bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  3. SIDE-BY-SIDE SUMMARY
# ══════════════════════════════════════════════════════════════
heading('3. Side-by-Side Comparison', 1)

img('comparison_v5_evolution', 'Traditional (1 chat) vs Structured Ideation (8 agents, parallel, evidence-backed)', 6.2)

table(
    ['', 'Traditional\n(1 AI chat)', 'This Method\n(8 agents)'],
    [
        ['You type',               '20+ messages',              '1 sentence + answer 10-15 questions'],
        ['Agents working',         '1 (does everything)',        '8 specialists (6 run in parallel)'],
        ['Internet research',      'No (uses training data)',   'Yes (every agent searches the web)'],
        ['Competitors found',      '0-2 (often wrong)',         '5+ with URLs, users, pricing'],
        ['Market data',            'Made up or vague',          'With source URLs you can click'],
        ['Critical analysis',      'None (AI agrees with you)', 'Dedicated evaluator + risk assessment'],
        ['Evidence per claim',     'None',                      'Source URL + confidence score'],
        ['Confidence score',       'No',                        'Yes (90%/60%/low per claim)'],
        ['Can you verify claims?', 'No',                        'Yes — every URL is clickable'],
        ['Output',                 'A chat transcript',         '13-section document with evidence'],
        ['Reproducible?',          'No',                        'Yes — same agents, same process'],
        ['Time',                   '~15 min',                   '~20 min (but 10x more output)'],
    ]
)

callout('💡 The key difference: traditional brainstorming gives you OPINIONS.\n'
        'This method gives you EVIDENCE. Every claim has a source you can check.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  4. HOW IT WORKS — Simple Explanation
# ══════════════════════════════════════════════════════════════
heading('4. How It Works (Simple Version)', 1)

body('Think of it like visiting a hospital. '
     'You don\'t get 1 doctor who does everything. '
     'You get a specialist for each concern:', bold=True)

table(
    ['Hospital Analogy', 'Our Pipeline', 'Why Separate?'],
    [
        ['Receptionist collects symptoms',   'Agent 1: asks you 10-15 questions',
         'Forces thorough info gathering\nbefore any diagnosis'],
        ['Lab runs blood tests',             'Agent 2-4: research market, tech,\ncompetitors (run simultaneously)',
         'Each looks at different angle\nwith fresh eyes'],
        ['Specialist reviews results',       'Agent 5-7: evaluate idea, UX, risks\n(run simultaneously)',
         'Evaluator hasn\'t seen your chat\nwith Agent 1 — no bias'],
        ['Doctor writes final report',       'Agent 8: compiles into 13-section spec',
         'Only compiles what specialists\nfound — adds nothing new'],
    ]
)

img('pipeline_v5_maxpower', 'The full pipeline: you type 1 prompt, 8 agents work for you, 6 run in parallel', 6.2)

heading('Why 8 Agents Instead of 1?', 2)
body('Three reasons:', bold=True)
body('1. Fresh perspective: Agent 5 (evaluator) has never seen your original idea. '
     'It only sees the research reports. So it evaluates objectively, '
     'like a jury that only sees evidence, not testimony.', bullet=True)
body('2. Parallel speed: Agents 2-4 run at the SAME TIME, not one after another. '
     'Same with Agents 5-7. So 8 agents finish in the time of 4.', bullet=True)
body('3. Quality control: The orchestrator checks each agent\'s output before '
     'passing it to the next phase. If an agent produces garbage, it gets re-run.', bullet=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  5. THE 8 AGENTS — What Each One Does
# ══════════════════════════════════════════════════════════════
heading('5. The 8 Agents — What Each One Does', 1)

img('agents_v5_detail', 'All 8 agents with their specialties and what they search for', 6.2)

agents_simple = [
    ('Phase 1: Understanding Your Idea', [
        ('🔍 Requirements Analyst',
         'Asks you 10-15 structured questions before doing anything else. '
         'Then challenges your assumptions: "You said no app does X — let me search to verify."',
         'A Requirements Brief with validated data'),
    ]),
    ('Phase 2: Researching (3 agents run at the same time)', [
        ('📊 Market Researcher',
         'Searches for: market size, demographics, trends, pricing benchmarks. '
         'Tests your market at extreme scales: "What if you had 1M users? What breaks?"',
         'Market Report with TAM/SAM/SOM + source URLs'),
        ('⚙️ Tech Scout',
         'Searches for APIs, frameworks, hosting costs. Checks real pricing pages. '
         'Finds patterns: "All the APIs you need have rate limits — you need a caching layer."',
         'Tech Report with API pricing + infrastructure costs'),
        ('🏢 Competitor Analyst',
         'Finds 5+ real competitors. Reads their App Store reviews. '
         'Generates creative alternatives: "What if your music app worked like a restaurant tasting menu?"',
         'Competitor Report with real URLs + creative ideas'),
    ]),
    ('Phase 3: Evaluating (3 agents run at the same time)', [
        ('⚖️ Idea Evaluator',
         'Scores your idea on 8 dimensions, each with evidence. '
         'Finds simplifications: "If we cut feature X, it also eliminates problems Y and Z."',
         'Evaluation Report with GO/PIVOT/KILL verdict'),
        ('🎨 UX Strategist',
         'Designs user journeys based on UX research from NNGroup/Baymard. '
         'Forces creative collisions: "What if your onboarding worked like a video game tutorial?"',
         'UX Strategy with personas + engagement hooks'),
        ('⚠️ Risk Assessor',
         'Identifies every risk, then asks: "How would I GUARANTEE this project fails?" '
         'Stress-tests riskat 1000x scale. Searches for legal requirements.',
         'Risk Assessment with mitigations + legal research'),
    ]),
    ('Phase 4: Final Document', [
        ('📋 Idea Crystallizer',
         'Takes all 7 reports and compiles them into 1 document. '
         'Does NOT add new info — only distills, connects, and traces every claim to its source.',
         '13-Section Idea Specification with confidence score'),
    ]),
]

for phase_name, agents in agents_simple:
    heading(phase_name, 2)
    for emoji_name, desc, output in agents:
        heading(emoji_name, 3)
        body(desc)
        body(f'Output: {output}', bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  6. WHAT YOU GET — The Output
# ══════════════════════════════════════════════════════════════
heading('6. What You Get: The 13-Section Idea Specification', 1)

body('After ~20 minutes, you receive a single document with:', bold=True)

table(
    ['Section', 'What It Contains', 'Why It Matters'],
    [
        ['1. Problem Statement', 'Pain point + who has it + how big', 'Validated — not assumed'],
        ['2. Target User', 'Persona + demographics + segment size', 'Based on real data'],
        ['3. Solution Overview', 'Core value + MVP features + USP', 'Differentiated from competitors'],
        ['4. Market Opportunity', 'TAM/SAM/SOM + growth rate + trends', 'Real numbers with URLs'],
        ['5. Competitive Landscape', '5+ real competitors + feature matrix', 'URLs you can click'],
        ['6. Technical Architecture', 'Stack + APIs + costs at scale', 'Pricing verified'],
        ['7. UX & User Journey', 'Onboarding + engagement + retention', 'Based on UX research'],
        ['8. Monetization Strategy', 'Revenue model + pricing benchmarks', 'Industry benchmarks cited'],
        ['9. Risk Analysis', 'Critical/high/medium/low risks', 'Legal requirements searched'],
        ['10. Evaluation Summary', '8-dimension scorecard + verdict', 'Evidence per score'],
        ['11. MVP Scope', 'Must-have / nice-to-have / out-of-scope', 'Realistic timeline'],
        ['12. Go-to-Market Sketch', 'Launch strategy + first 100 users', 'Based on competitor analysis'],
        ['13. Evidence Trail', 'ALL source URLs + confidence scores', 'Full audit trail'],
    ]
)

heading('Confidence Score — How Much Can You Trust This?', 2)
body('Every claim in the document has a confidence level:', bold=True)
table(
    ['Level', 'What It Means', 'Example'],
    [
        ['✅ Validated (90%+)', '3+ sources agree', 'Spotify has 600M+ users (Wikipedia, TechCrunch, Statista)'],
        ['🟡 Likely (60-89%)', '1-2 sources found', 'Indie music market ~$4B (Grand View Research)'],
        ['🟠 Low confidence', 'Reasoning only', 'Estimated 5% conversion based on similar apps'],
        ['❌ Unvalidated', 'No data found', 'No data on Vietnamese indie music app market [NEEDS RESEARCH]'],
    ]
)
body('The final document includes an Overall Confidence Score — '
     'so you know exactly how much of the spec is backed by data vs. assumptions.', bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  7. HOW TO USE IT — Step by Step
# ══════════════════════════════════════════════════════════════
heading('7. How to Use It (Step by Step)', 1)

img('setup_v5', 'One-time setup: copy files, then use forever', 5.5)

heading('One-Time Setup (5 minutes)', 2)
body('1. Copy the 9 agent files to your project:', bold=True)
code('.cursor/\n'
     '  agents/                    <-- Cursor auto-detects these\n'
     '    ideation-orchestrator.md\n'
     '    requirements-analyst.md\n'
     '    market-researcher.md\n'
     '    tech-scout.md\n'
     '    competitor-analyst.md\n'
     '    idea-evaluator.md\n'
     '    ux-strategist.md\n'
     '    risk-assessor.md\n'
     '    idea-crystallizer.md\n'
     '  mcp.json                   <-- MCP servers (optional)\n')

body('2. That\'s it. No installation, no packages, no API keys.', bold=True)

heading('Every Time You Have a New Idea', 2)
body('Step 1: Open Cursor → Agent Mode', bullet=True)
body('Step 2: Type your idea (even 1 sentence is enough):', bullet=True)
callout('💬 /ideation-orchestrator "I want to make a music app that helps discover indie Vietnamese artists"')
body('Step 3: Answer 10-15 questions from the Requirements Analyst', bullet=True)
body('Step 4: Wait ~15 minutes while 6 agents research in parallel', bullet=True)
body('Step 5: Receive your 13-section Idea Specification', bullet=True)

heading('Example: What Your Interaction Looks Like', 2)
body('The Requirements Analyst will ask questions like:')
table(
    ['Category', 'Example Question'],
    [
        ['WHO', '"Who is your target user? Age range? Are they tech-savvy?"'],
        ['WHAT', '"What are the 3 core features? Mobile or web? Offline needed?"'],
        ['WHY', '"What apps exist that you don\'t like? Why?"'],
        ['HOW', '"What\'s your budget for an MVP? Timeline?"'],
        ['CONSTRAINTS', '"Any licensing concerns? Must use specific technology?"'],
    ]
)
body('You answer briefly. The agent follows up on vague answers. '
     'Then all 7 other agents work automatically.', bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  8. WHY THIS WORKS — The Techniques
# ══════════════════════════════════════════════════════════════
heading('8. What Makes It Smart: Built-In Thinking Techniques', 1)

body('Each agent doesn\'t just search — it THINKS using proven problem-solving techniques:')

img('techniques_v5', 'Each agent uses a specific thinking technique to go deeper than simple search', 6.2)

table(
    ['Technique', 'What It Does', 'Example'],
    [
        ['Assumption Flipping', 'Challenges what you believe is true.\n"What if the opposite were true?"',
         'You say "users want more features"\n→ Flip: "What if they want FEWER?"\n→ Insight: Focus on simplicity wins'],
        ['Extreme Scale Testing', 'Tests your idea at 1000x and 0.001x.\n"What breaks? What survives?"',
         'At 1M users: API costs = $50K/mo\nAt 100 users: does app make sense?\n→ Need cost alerts + min viable audience'],
        ['Creative Collision', 'Forces unrelated concepts together.\n"What if music app × restaurant?"',
         '→ "Daily tasting flight of 5 curated songs"\n→ No algorithm, no choice paralysis\n→ Something NO competitor has'],
        ['Pattern Recognition', 'Spots patterns across 3+ technologies.\n"What\' universal here?"',
         'All APIs rate-limit at scale\n→ Need caching layer for ALL APIs\n→ Architecture decision made early'],
        ['Simplification', 'Finds 1 change that eliminates\nmultiple problems at once.',
         '"Don\'t host music files at all"\n→ Eliminates: licensing, CDN costs,\ncatalog management, AND moderation'],
    ]
)

body('These techniques come from ClaudeKit (open-source by Anthropic) — '
     'proven problem-solving frameworks adapted for AI agents.', bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  9. BUILT-IN SAFETY — Quality Control
# ══════════════════════════════════════════════════════════════
heading('9. Built-In Safety: Quality Gates & When-Stuck Protocol', 1)

heading('Quality Gates', 2)
body('The orchestrator doesn\'t blindly pass data between phases. '
     'After Phase 2 and Phase 3, it checks:', bold=True)
body('✓ Does each report have at least 5 source URLs?', bullet=True)
body('✓ Are all required sections filled (no empty headings)?', bullet=True)
body('✓ Is there at least 1 data point per major claim?', bullet=True)
body('If an agent fails → it gets re-run with better instructions.', bold=True)

heading('When-Stuck: What If the Agent Can\'t Find Data?', 2)
body('Real problem: some searches return nothing. Every agent has a fallback:')
table(
    ['Step', 'What the Agent Does'],
    [
        ['1. Broaden', 'Remove qualifiers, try synonyms, search broader category'],
        ['2. Alternative sources', 'Try Reddit, HackerNews, Quora, industry reports'],
        ['3. Adjacent markets', 'Search parent/sibling categories instead'],
        ['4. Accept the gap', 'Mark as [NO DATA] — NEVER fabricate an answer'],
    ]
)

callout('🔒 Rule: An agent will NEVER make up data to fill a gap.\n'
        'It will always tell you "I couldn\'t find this" rather than invent numbers.')

# ══════════════════════════════════════════════════════════════
#  10. ALSO WORKS WITH
# ══════════════════════════════════════════════════════════════
heading('10. Also Works With Other AI Tools', 1)
body('This method is designed for Cursor, but the principles work anywhere:')
table(
    ['Tool', 'How to Use'],
    [
        ['Cursor IDE', 'Copy agent files to .cursor/agents/ (recommended — native integration)'],
        ['ChatGPT', 'Run 4 separate conversations, paste agent prompt as first message each time'],
        ['Claude', 'Use Projects feature — 1 project per agent, paste outputs between them'],
        ['Windsurf', 'Use .windsurfrules files (similar to Cursor)'],
        ['Any AI', 'The core idea: separate sessions, separate roles, documents as handoff'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  11. TIPS & COMMON MISTAKES
# ══════════════════════════════════════════════════════════════
heading('11. Tips & Common Mistakes', 1)

heading('Do', 2)
for tip in [
    'Keep the seed SHORT (even 1 sentence). The first agent\'s job is to expand it.',
    'Answer questions honestly — short answers are fine, the agent follows up.',
    'Read the confidence scores — low-confidence claims need manual verification.',
    'Challenge Agent 5\'s scoring if something feels wrong: "Why did you give this a 4?"',
    'Use the SCAMPER variants — they often produce the winning idea.',
    'This Idea Spec is the INPUT for Practice 2 (Idea → Plan).',
]:
    body(tip, bullet=True)

heading('Don\'t', 2)
for tip in [
    'Don\'t skip the questioning phase — bad requirements = bad everything.',
    'Don\'t accept scores without reading the evidence.',
    'Don\'t modify the final spec — if something is wrong, re-run the responsible agent.',
    'Don\'t use this for trivial ideas — if you already know exactly what to build, skip to Practice 2.',
]:
    body(tip, bullet=True)

# ══════════════════════════════════════════════════════════════
#  12. WHAT'S NEXT
# ══════════════════════════════════════════════════════════════
heading('12. What\'s Next: Practice 2 & 3', 1)
body('This is Practice 1 of a 3-part series. Each practice feeds into the next:', bold=True)
table(
    ['Practice', 'Name', 'Input', 'Output'],
    [
        ['1 (this)', 'Multi-Agent Ideation', '1 sentence idea', '13-section Idea Specification'],
        ['2 (next)', 'Idea to Plan', 'Idea Specification', 'Architecture + user stories + timeline'],
        ['3 (last)', 'Plan to Code', 'Implementation Plan', 'Working code + tests'],
    ]
)
body('The Idea Specification from Practice 1 becomes the input for Practice 2. '
     'Different agents, same principle: separate specialists, documents as handoff, '
     'evidence over opinions.', bold=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  SOURCES
# ══════════════════════════════════════════════════════════════
heading('Sources & References', 1)
body('This method combines ideas from:', bold=True)

img('sources_v5', 'The open-source frameworks that power this method', 5.5)

table(
    ['Source', 'What We Used', 'URL'],
    [
        ['BMAD Method', 'Multi-agent personas, orchestrator pattern',
         'github.com/bmad-code-org/BMAD-METHOD'],
        ['ClaudeKit', '5 thinking techniques + context engineering',
         'github.com/anthropics/claudekit'],
        ['Cursor Subagents', 'Native parallel execution + context isolation',
         'cursor.com/learn/agents'],
        ['context7.com', 'API documentation via llms.txt protocol',
         'context7.com'],
    ]
)

# FOOTER
doc.add_paragraph()
ft = doc.add_paragraph(); ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ft.add_run('Report generated: 2026-04-16  |  Practice 1: Multi-Agent Ideation  |  Video Editor')
r.font.size=Pt(8); r.font.color.rgb=RGBColor(0x99,0x99,0x99); r.font.name='Calibri'

out = r'd:\PROJECT\MyMusic\docs\practice1\Practice-1-Multi-Agent-Ideation.docx'
doc.save(out)
print(f'[OK] Word document saved to: {out}')
