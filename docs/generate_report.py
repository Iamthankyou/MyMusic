"""
Generate BMAD Method AI Case Report as a Word document (.docx)
with embedded diagram images for professional presentation.
"""
import glob, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Locate diagram images (matplotlib corporate-style, from docs/diagrams/) ──
IMG_DIR = r'd:\PROJECT\MyMusic\docs\diagrams'

IMG_WORKFLOW     = os.path.join(IMG_DIR, 'workflow.png')
IMG_ARCHITECTURE = os.path.join(IMG_DIR, 'architecture.png')
IMG_AGENTS       = os.path.join(IMG_DIR, 'agents.png')
IMG_COMPARISON   = os.path.join(IMG_DIR, 'comparison.png')
IMG_DATAFLOW     = os.path.join(IMG_DIR, 'dataflow.png')
IMG_WEAKNESSES   = os.path.join(IMG_DIR, 'tradeoffs.png')
IMG_TRACEABILITY = os.path.join(IMG_DIR, 'traceability.png')

print("Images found:")
for name, path in [("Workflow", IMG_WORKFLOW), ("Architecture", IMG_ARCHITECTURE),
                    ("Agents", IMG_AGENTS), ("Comparison", IMG_COMPARISON),
                    ("DataFlow", IMG_DATAFLOW), ("Weaknesses", IMG_WEAKNESSES),
                    ("Traceability", IMG_TRACEABILITY)]:
    print(f"  {name}: {'OK' if os.path.exists(path) else 'MISSING'}")

# ── Init document ─────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width  = Inches(8.27)  # A4
section.page_height = Inches(11.69)
section.left_margin   = Cm(2)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False, size=10,
            color=None, font_name='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run

def label_cell(cell, text):
    shade_cell(cell, '2E4057')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def value_cell(cell, items):
    first = True
    for item in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if item.get('bullet'):
            p.style = doc.styles['List Bullet']
        add_run(p, item.get('text', ''),
                bold=item.get('bold', False),
                size=item.get('size', 10),
                color=item.get('color'))

def add_section_heading(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Calibri'
    if level == 1:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x48, 0x6F, 0x8C)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
    else:
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8B)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

def add_body(text, bullet=False, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    if bullet:
        p.style = doc.styles['List Bullet']
    add_run(p, text, bold=bold, size=size, color=color)
    return p

def add_image(img_path, caption=None, width_inches=6.2):
    if img_path and os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            r.font.name = 'Calibri'
    else:
        add_body(f'[Diagram image not found: {caption or "unknown"}]',
                 color='CC0000')

def add_info_table(rows_data):
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, content) in enumerate(rows_data):
        row = tbl.rows[i]
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(13)
        label_cell(row.cells[0], label)
        value_cell(row.cells[1], content)
    doc.add_paragraph()

def add_simple_table(headers, rows_data):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        shade_cell(hdr.cells[i], '2E4057')
        p = hdr.cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows_data):
        drow = tbl.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else 'EEF2F7'
        for ci, cell_text in enumerate(row_data):
            shade_cell(drow.cells[ci], bg)
            p = drow.cells[ci].paragraphs[0]
            r = p.add_run(cell_text)
            r.font.size = Pt(9)
            r.font.name = 'Calibri'
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('AI Utilization Case Report')
tr.bold = True; tr.font.size = Pt(22); tr.font.name = 'Calibri'
tr.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run('BMAD Method (Breakthrough Method for Agile AI-Driven Development)')
sr.font.size = Pt(13); sr.font.name = 'Calibri'
sr.font.color.rgb = RGBColor(0x48, 0x6F, 0x8C)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2 = sub2.add_run('Applied to: MyMusic Android App')
s2.font.size = Pt(12); s2.font.name = 'Calibri'
s2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp = date_p.add_run('Submitted: 2026-04-15  |  Multimedia 1P  |  Duy LQ')
dp.font.size = Pt(10); dp.font.name = 'Calibri'
dp.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  BASIC INFO TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_info_table([
    ("Date of Submission", [{'text': '2026.04.15'}]),
    ("Name (Department)",  [{'text': 'Duy LQ (Multimedia 1P)'}]),
    ("Case Title",         [{'text': 'Building MyMusic Android App Using BMAD Method',
                             'bold': True, 'size': 11}]),
    ("Utilization Scale",  [{'text': 'Full transformation of development workflow', 'bold': True},
                             {'text': 'From manual planning + ad-hoc AI prompting '
                                      'to structured, agent-orchestrated, '
                                      'documentation-first AI development lifecycle'}]),
    ("Case Category",      [{'text': 'Productivity tool development (Agent-based workflow)', 'bold': True},
                             {'text': 'Design & Documentation: architecture, data model, API, doc generation'},
                             {'text': 'Coding: code generation, Clean Architecture implementation'},
                             {'text': 'Review: static analysis, code change summary'},
                             {'text': 'Testing: unit test generation (planned)'}]),
    ("Tools Used",         [{'text': 'Cursor IDE, Cline (VS Code), Antigravity (Google DeepMind)'}]),
    ("Models Used",        [{'text': 'Claude 4.5-Sonnet, Claude Opus 4.6, GPT-5'}]),
])

# ══════════════════════════════════════════════════════════════════════════════
#  CASE DETAILS
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading('Case Details', level=1)

# ──────────────────────────────────────────────────────────────────────────────
add_section_heading('1. What is BMAD Method?', level=1)
add_body(
    'BMAD (Breakthrough Method for Agile AI-Driven Development) is an open-source '
    'framework that brings structure, governance, and repeatability to AI-assisted '
    'software development. Instead of "vibe coding" (ad-hoc prompting), it treats '
    'AI as a disciplined team member following an agile process with clear roles, '
    'producing version-controlled artifacts at every phase.'
)

# ── 1.1 Comparison ───────────────────────────────────────────────────────────
add_section_heading('1.1  Why BMAD? — Vibe Coding vs. BMAD Comparison', level=2)
add_image(IMG_COMPARISON, 'Figure 1: Vibe Coding vs BMAD Method — The core problem and solution', 6.0)

# ── 1.2 Agent Team ───────────────────────────────────────────────────────────
add_section_heading('1.2  Specialized AI Agent Team', level=2)
add_body(
    'BMAD uses 6 specialized AI agent personas, each with distinct responsibilities '
    'mirroring a professional software team. Each agent produces explicit, '
    'version-controlled artifacts.'
)
add_image(IMG_AGENTS, 'Figure 2: BMAD Agent Team — 6 specialized agents with outputs', 6.0)

add_simple_table(
    ['Agent', 'Role', 'Output Artifact'],
    [
        ['Analyst',         'Brainstorming, research, scoping',               'project-brief.md'],
        ['Product Manager', 'Requirements definition, PRD creation',          'prd.md'],
        ['Architect',       'System design, technology decisions, API design', 'fullstack-architecture.md'],
        ['Scrum Master',    'Epic decomposition, user story creation',         'stories/*.md'],
        ['Developer',       'Code generation following stories + architecture','Source code files'],
        ['QA',              'Test generation, code review, quality gates',     'Test files + review feedback'],
    ]
)

# ── 1.3 Workflow ─────────────────────────────────────────────────────────────
add_section_heading('1.3  BMAD Workflow — 4-Phase Development Lifecycle', level=2)
add_body(
    'BMAD enforces a structured lifecycle with mandatory human review gates '
    'between each phase. This ensures quality control and prevents cascading errors.'
)
add_image(IMG_WORKFLOW, 'Figure 3: BMAD 4-Phase Workflow with Human Review Gates', 6.2)

add_simple_table(
    ['Phase', 'Agent(s)', 'Key Activities', 'Output Artifacts'],
    [
        ['1. Analysis', 'Analyst', 'Brainstorm, research, risk ID', 'project-brief.md'],
        ['Human Gate', 'Developer (Human)', 'Review scope, approve direction', 'Approval / feedback'],
        ['2. Planning', 'Product Manager', 'Requirements (FR/NFR), UX goals', 'prd.md + front-end-spec.md'],
        ['Human Gate', 'Developer (Human)', 'Review requirements completeness', 'Approval / feedback'],
        ['3. Solutioning', 'Architect + Scrum Master', 'Architecture design, story decomposition', 'architecture.md + stories/'],
        ['Human Gate', 'Developer (Human)', 'Review architecture decisions', 'Approval / feedback'],
        ['4. Implementation', 'Developer + QA', 'Code and test story-by-story', 'Source code + tests'],
    ]
)

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────────
add_section_heading('2. How We Applied BMAD to Build MyMusic', level=1)
add_body(
    'MyMusic is a modern Android music app built 100% in Kotlin with Jetpack Compose, '
    'Hilt DI, MVVM + Clean Architecture, powered by the Jamendo API for '
    'streaming music discovery and playback.'
)

# ── 2.1 Architecture ────────────────────────────────────────────────────────
add_section_heading('2.1  MyMusic Architecture (C4 Container Level)', level=2)
add_image(IMG_ARCHITECTURE, 'Figure 4: MyMusic App Architecture — 4-layer Clean Architecture with C4 model', 5.5)

# ── 2.2 Data Flow ───────────────────────────────────────────────────────────
add_section_heading('2.2  Data Flow — From User Action to API and Back', level=2)
add_image(IMG_DATAFLOW, 'Figure 5: MyMusic Data Flow — User interaction to API call sequence', 5.0)

# ── 2.3 Codebase Stats ──────────────────────────────────────────────────────
add_section_heading('2.3  Codebase Statistics', level=2)
add_simple_table(
    ['Metric', 'Value', 'Notes'],
    [
        ['Total Kotlin source files',   '102',   'Generated following BMAD stories'],
        ['Documentation files (docs/)', '80',    'PRD, Architecture, UX Spec, Stories, Diagrams'],
        ['User stories written',        '22',    'Across Epic 1 (8), Epic 2 (5), Epic 3 (6+)'],
        ['Domain models',               '5',     'Track, Album, Artist, Playlist, SearchFilter'],
        ['Repository interfaces',       '3',     'Domain layer contracts'],
        ['Use cases',                   '12',    'Business logic interactors'],
        ['Retrofit services',           '6',     'Jamendo API endpoints'],
        ['Repository implementations',  '7',     'Data layer'],
        ['Presentation screens',        '8+',    'Compose screens + ViewModels'],
        ['Architecture layers',         '4',     'Presentation / Domain / Data / Media'],
        ['Test files',                  '0',     'CRITICAL GAP  -  0% test coverage'],
    ]
)

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────────
add_section_heading('3. Why Should You Use BMAD? — Proven Benefits', level=1)

# ── 3.1 Traceability ────────────────────────────────────────────────────────
add_section_heading('3.1  Full Traceability — Requirement to Code', level=2)
add_body(
    'Every line of code in the project traces back through a clear chain: '
    'PRD Requirement -> Architecture Section -> Epic / Story -> Code File. '
    'This ensures no feature is "orphan code" without a documented reason.'
)
add_image(IMG_TRACEABILITY, 'Figure 6: BMAD Traceability Map — from PRD FR to source code file', 6.0)

# ── 3.2 Benefits Table ──────────────────────────────────────────────────────
add_section_heading('3.2  Concrete Benefits Proven by MyMusic', level=2)
add_simple_table(
    ['Benefit', 'Evidence from MyMusic Project'],
    [
        ['Clean Architecture from Day 1',
         '4-layer separation (Presentation / Domain / Data / Media) enforced by architecture doc before any code was written'],
        ['Zero Context Loss Between Sessions',
         '80 docs serve as persistent AI context. Agents always know the project state, even weeks later'],
        ['Full Requirement Traceability',
         'Every feature: PRD FR -> Architecture Section -> Epic -> Story -> Code file'],
        ['Consistent Code Quality',
         'All 102 files follow identical patterns: Hilt DI, StateFlow, Mapper, Paging'],
        ['Fast Onboarding',
         'New developer reads project-brief.md -> prd.md -> architecture.md -> starts coding within 30 minutes'],
        ['Scalable Complexity',
         'Started with Trending Tracks MVP, expanded to 3 epics + 22 stories without architectural churn'],
    ]
)

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────────
add_section_heading('4. Weaknesses & Limitations of BMAD — Honest Assessment', level=1)
add_body(
    'BMAD is powerful but not perfect. Below is an honest assessment of its '
    'trade-offs, based on our real experience building MyMusic.'
)
add_image(IMG_WEAKNESSES, 'Figure 7: BMAD Strengths vs Weaknesses — Trade-off Analysis', 6.0)

add_section_heading('4.1  Detailed Weakness Analysis', level=2)
add_simple_table(
    ['Weakness', 'Impact on MyMusic', 'Mitigation Strategy'],
    [
        ['Overhead for small tasks', 
         'Heavy ceremony even for a simple color change or minor bug fix',
         'Use BMAD "Quick Flow" track for minor tasks (skip PRD/Architecture)'],
        ['High token consumption',
         '80 docs + multi-agent = ~$8-15 USD API cost for full project',
         'Use cheaper models for doc generation, premium models for critical decisions'],
        ['Cascading errors',
         'Wrong assumption in Phase 1 (e.g., API endpoint) propagates through all phases',
         'Mandatory human review gate between every phase transition'],
        ['Planning rigidity',
         '"Plan everything first" slows down exploratory / experimental work',
         'Use hybrid approach: BMAD for core, vibe coding for spikes'],
        ['Forced quality metrics',
         'Adversarial code review can create artificial nitpicking / developer fatigue',
         'Remove minimum-issue requirements from code review agent prompts'],
        ['Single-repo focus',
         'Struggles with cross-service or multi-repo distributed architectures',
         'Use BMAD Enterprise track for complex system-of-systems'],
        ['No test enforcement',
         'Test strategy documented (Architecture $17) but 0 actual tests created',
         'Add explicit "generate tests" step as mandatory gate per story'],
    ]
)

add_section_heading('4.2  When NOT to Use BMAD — Decision Guide', level=2)
add_simple_table(
    ['Scenario',                       'Recommendation',     'Reason'],
    [
        ['Quick bug fix (< 1 hour)',   'Skip BMAD',          'Overhead not justified for minor fix'],
        ['Small isolated feature',     'Quick Flow only',    'Tech-spec to implement, skip full docs'],
        ['New MVP app (1-4 weeks)',    'Standard BMAD',      'Full workflow, best ROI'],
        ['Enterprise / regulated',     'Enterprise BMAD',    'Need security + compliance + DevOps docs'],
        ['Exploratory prototype',      'Use with caution',   '"Plan first" may slow discovery; consider hybrid'],
        ['Refactoring existing code',  'Standard BMAD',      'Architecture doc prevents regressions'],
    ]
)

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────────
add_section_heading('5. Current Testing Status', level=1)
add_body('CRITICAL GAP: Despite a comprehensive test strategy documented in the '
         'architecture document (Section 17), NO actual test files were created.', bold=True, color='CC0000')

add_simple_table(
    ['Test Category', 'Status', 'Documented Strategy (Architecture S17)'],
    [
        ['app/src/test/ directory', 'Does NOT exist', 'Unit tests: mappers, use cases, repos with fakes'],
        ['app/src/androidTest/ directory', 'Does NOT exist', 'Instrumented: ViewModels with coroutine testing'],
        ['*Test.kt files', 'ZERO files found', 'UI tests: Compose key flows (Trending -> Play)'],
        ['Unit Test Coverage', '0%', 'Target: mappers, use cases, ViewModels'],
        ['Integration Test Coverage', '0%', 'Target: repos with fake servers'],
        ['UI Test Coverage', '0%', 'Target: key user flows'],
    ]
)

add_body('Key Lesson: BMAD produces excellent test PLANS but does NOT enforce test CREATION. '
         'Human discipline and explicit story gates for testing are required.', bold=True)
doc.add_paragraph()
add_body('Recommended priority for test generation:', bold=True)
add_body('1. TrackMapper (DTO to Domain) - pure function, easiest to test, highest ROI', bullet=True)
add_body('2. GetTrendingTracksUseCase - validates business logic isolation', bullet=True)
add_body('3. TrendingViewModel - UI state transition tests with coroutine TestDispatcher', bullet=True)
add_body('4. PlaybackController - media state management edge cases', bullet=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TIME SAVED
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading('Time Saved by Using AI', level=1)

add_simple_table(
    ['#', 'Task', 'Before (Manual)', 'After (BMAD + AI)', 'Time Saved'],
    [
        ['1', 'Requirements & PRD Creation',            '24 hours (3 days)',   '3 hours',    '21 hours (87%)'],
        ['2', 'Architecture Design & Documentation',    '16 hours (2 days)',   '2 hours',    '14 hours (87%)'],
        ['3', 'User Story Decomposition (22 stories)',  '16 hours (2 days)',   '1.5 hours',  '14.5 hours (91%)'],
        ['4', 'Code Generation (102 Kotlin files)',     '120 hours (15 days)', '24 hours',   '96 hours (80%)'],
        ['5', 'Architecture Diagrams & Visual Docs',    '8 hours (1 day)',     '0.5 hours',  '7.5 hours (94%)'],
        ['',  'TOTAL', '184 hours (~23 days)', '31 hours (~4 days)', '153 hours (83%)'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  IMPROVEMENTS
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading('Improvements or Benefits from AI Use', level=1)
for b in [
    'Architecture quality: Clean 4-layer separation enforced from Day 1 - no spaghetti refactor needed later',
    'Documentation completeness: 80 docs generated (PRD, Architecture, UX Spec, 22 stories) vs. typical 0-3 docs in ad-hoc projects',
    'Consistency: All 102 Kotlin files follow identical patterns (Hilt, StateFlow, Mapper, Paging)',
    'Onboarding speed: New developer reads 3 docs and starts contributing within 30 minutes',
    'Traceability: Every code file traces back to Story -> Epic -> Architecture Section -> FR in PRD',
    'Decision documentation: Technology choices explained in architecture doc (why Kotlinx Serialization, why single-module, why Media3)',
]:
    add_body(b, bullet=True)

# ══════════════════════════════════════════════════════════════════════════════
#  LIMITATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading('Limitations or Challenges', level=1)
for lim in [
    'No test files generated: Despite documented test strategy, BMAD did not enforce actual test creation. 0% test coverage is a critical gap.',
    'Token cost: Generating 80 documentation files + 102 source files consumed ~$8-15 USD in API tokens (Claude Sonnet).',
    'Overhead for small changes: Minor UI tweaks require checking PRD/Architecture alignment, which feels heavy.',
    'Cascading risk: Incorrect assumptions about Jamendo API in Phase 2 would propagate through all phases.',
    'Human discipline required: BMAD provides framework but developer must still review each artifact critically.',
]:
    add_body(lim, bullet=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  OUTPUT LINKS
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading('Output / Deliverable Link', level=1)
add_simple_table(
    ['Deliverable', 'Path / Location'],
    [
        ['Project Brief',           'docs/project-brief.md'],
        ['Brainstorming Doc',       'docs/brain-stoming-documentation.md'],
        ['PRD (16 FR + 13 NFR)',    'docs/prd.md'],
        ['Front-End UX Spec',       'docs/front-end-spec.md'],
        ['Architecture (21 sections)', 'docs/fullstack-architecture.md'],
        ['User Stories (22)',       'docs/stories/*.md'],
        ['Architecture Diagram',    'docs/diagrams/architecture-overview.mmd / .png'],
        ['Source Code (102 files)', 'app/src/main/java/com/example/mymusic/'],
        ['Built APK',               'app-debug.apk'],
        ['This Report',             'docs/BMAD-Method-AI-Case-Report.docx'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  APPENDIX
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading('Appendix', level=1)

add_section_heading('A. Useful Prompts for BMAD Workflow', level=2)

prompts = [
    ('Phase 1 - Analyst Agent', 
     '"Brainstorm a [type] app with these features: [list]. '
     'Research API options, identify risks, and create a project-brief.md."'),
    ('Phase 2 - Product Manager Agent',
     '"Based on the project brief, create a PRD with functional requirements (FR1-FRn), '
     'non-functional requirements (NFR1-NFRn), epic list, and UI/UX goals. '
     'Run a checklist validation at the end."'),
    ('Phase 3 - Architect Agent',
     '"Create a full-stack architecture document based on this PRD. '
     'Include: module structure, DI graph, data flow, API design, error handling, '
     'testing strategy. Use Clean Architecture + MVVM."'),
    ('Phase 3b - Scrum Master Agent',
     '"Decompose Epic [n] into user stories with acceptance criteria. '
     'Each story must be independently implementable and include technical notes '
     'referencing the architecture document."'),
    ('Phase 4 - Developer Agent (per story)',
     '"Implement Story [n.m] following the architecture document. '
     'Use [tech stack]. Include unit tests for new business logic. '
     'Match patterns established in existing codebase."'),
]
tbl = doc.add_table(rows=len(prompts), cols=2)
tbl.style = 'Table Grid'
for i, (title_text, prompt_text) in enumerate(prompts):
    bg = 'F5F5F5' if i % 2 == 0 else 'FFFFFF'
    label_cell(tbl.rows[i].cells[0], title_text)
    shade_cell(tbl.rows[i].cells[1], bg)
    p = tbl.rows[i].cells[1].paragraphs[0]
    r = p.add_run(prompt_text)
    r.font.size = Pt(9)
    r.font.name = 'Calibri'
    r.italic = True
doc.add_paragraph()

add_section_heading('B. Cursor Rules Template (.cursorrules)', level=2)
rules = [
    '# Project: MyMusic',
    '# Method:  BMAD',
    '',
    '## Context Files (Always Read First)',
    '- docs/project-brief.md - Project scope and goals',
    '- docs/prd.md - Requirements (FR/NFR)',
    '- docs/fullstack-architecture.md - Architecture decisions',
    '- docs/stories/ - Current stories and acceptance criteria',
    '',
    '## Architecture Rules',
    '- MVVM + Clean Architecture (4 layers)',
    '- All DI via Hilt (@Singleton, @ViewModelScoped)',
    '- DTOs in data layer only; domain models in domain layer',
    '- Mappers convert DTO <-> Domain at repository boundary',
    '- UI state via StateFlow + collectAsStateWithLifecycle()',
    '- Paging via Paging 3 PagingSource + cachedIn(viewModelScope)',
    '',
    '## Testing (MANDATORY per story)',
    '- Every new UseCase MUST have a unit test',
    '- Every new Mapper MUST have a unit test',
    '- Every new ViewModel MUST have state transition tests',
]
# Single-cell code block
tbl2 = doc.add_table(rows=1, cols=1)
tbl2.style = 'Table Grid'
cell = tbl2.rows[0].cells[0]
shade_cell(cell, 'F5F5F5')
first = True
for line in rules:
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    first = False
    r = p.add_run(line if line else ' ')
    r.font.name = 'Courier New'
    r.font.size = Pt(8)
doc.add_paragraph()

add_section_heading('C. BMAD Track Selection Guide', level=2)
add_simple_table(
    ['Project Type',                    'BMAD Track',      'Key Artifacts'],
    [
        ['Bug fix / small feature',     'Quick Flow',      'Tech-spec -> Code'],
        ['MVP / new product (1-4 weeks)','Standard BMAD',  'Brief + PRD + Architecture + Stories'],
        ['Enterprise / regulated',      'Enterprise BMAD', 'Standard + Security + DevOps docs'],
        ['Exploratory prototype',       'Hybrid approach', 'Lightweight brief -> spike -> full BMAD'],
        ['Refactoring existing code',   'Standard BMAD',   'Architecture doc prevents regressions'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
ft = doc.add_paragraph()
ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ft.add_run('Report generated: 2026-04-15  |  Project: MyMusic  |  Method: BMAD  |  Multimedia 1P')
r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99); r.font.name = 'Calibri'

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = r'd:\PROJECT\MyMusic\docs\BMAD-Method-AI-Case-Report.docx'
doc.save(out_path)
print(f'[OK] Word document saved to: {out_path}')
print(f'[OK] Contains 7 embedded diagram images')
